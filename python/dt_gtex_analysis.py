import os
base_dir = os.path.dirname(os.path.abspath(__file__))
"""
DT vs GTEx analysis — mirrors predict_gtex.py but with Decision Tree.
Generates per-cohort:
  - feature_importance_DT_TCGA_<cohort>_vs_gtex.png
  - pathway_DT_TCGA_<cohort>_vs_gtex.png
  - confusion_matrix_DT_TCGA_<cohort>_vs_GTEx.png (re-generated with 5-fold)
  - DT_GTEx_Summary.csv
Then embeds all figures into the docx after Fig 81.
"""
import os, csv, time, glob, requests
import numpy as np
import pandas as pd
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import seaborn as sns
from sklearn.tree import DecisionTreeClassifier
from sklearn.model_selection import StratifiedKFold
from sklearn.metrics import (classification_report, confusion_matrix, f1_score)
import docx
from docx.oxml.ns import qn
from docx.shared import Inches

BASE      = os.path.dirname(base_dir)
DS_DIR    = BASE + "/python/ml_dataset_project_batches"
OUT_DIR   = BASE + "/outputs"
FIG_DIR   = OUT_DIR + "/figures"
DOCX_PATH = BASE + "/Bozza Identification of Biomarkers from Gene Expression Data.docx"

TOP_N = 50
FOLDS = 5
SEED  = 42

COHORT_LABELS = {
    "TCGA_BRCA": "Breast Cancer (BRCA)",
    "TCGA_LUAD": "Lung Adenocarcinoma (LUAD)",
    "TCGA_LUSC": "Lung Squamous Cell Carcinoma (LUSC)",
    "TCGA_OV":   "Ovarian Cancer (OV)",
    "TCGA_PRAD": "Prostate Adenocarcinoma (PRAD)",
    "TCGA_STAD": "Stomach Adenocarcinoma (STAD)",
    "TCGA_THCA": "Thyroid Carcinoma (THCA)",
}
TISSUE_MAP = {
    "TCGA_BRCA": "Breast - Mammary Tissue",
    "TCGA_LUAD": "Lung",
    "TCGA_LUSC": "Lung",
    "TCGA_OV":   "Ovary",
    "TCGA_PRAD": "Prostate",
    "TCGA_STAD": "Stomach",
    "TCGA_THCA": "Thyroid",
}

def load_symbol_map():
    named = OUT_DIR + "/Biomarkers_Final_Named.csv"
    if os.path.exists(named):
        df = pd.read_csv(named)[["Gene","symbol"]].dropna()
        return dict(zip(df["Gene"], df["symbol"]))
    return {}

ENRICHR_URL = "https://maayanlab.cloud/Enrichr"

def enrichr_kegg(gene_list, label):
    try:
        r = requests.post(ENRICHR_URL + "/addList",
                          files={"list": (None, "\n".join(gene_list)),
                                 "description": (None, label)},
                          timeout=30)
        if r.status_code != 200:
            return None
        uid = r.json()["userListId"]
        time.sleep(1.5)
        r2 = requests.get(ENRICHR_URL + "/enrich",
                          params={"userListId": uid,
                                  "backgroundType": "KEGG_2021_Human"},
                          timeout=30)
        if r2.status_code != 200:
            return None
        return [{"term": x[1], "adj_p": x[6]} for x in r2.json().get("KEGG_2021_Human", [])[:10]]
    except Exception as e:
        print("  Enrichr error:", e)
        return None

def plot_kegg(rows, cohort, out_path):
    if not rows:
        return False
    terms = [r["term"] for r in rows]
    pvals = [-np.log10(r["adj_p"] + 1e-10) for r in rows]
    fig, ax = plt.subplots(figsize=(9, 5))
    ax.barh(terms[::-1], pvals[::-1], color="#2196F3")
    ax.set_xlabel("-log10(adj. p-value)")
    ax.set_title("%s — DT KEGG pathway enrichment (top-10)" % cohort)
    plt.tight_layout()
    plt.savefig(out_path, dpi=150, bbox_inches="tight")
    plt.close()
    return True

def plot_importance(imp, genes, sym_map, label, out_path):
    names = [sym_map.get(g, g) for g in genes]
    idx   = np.argsort(imp)
    colors = plt.cm.RdYlBu_r(np.linspace(0.05, 0.95, len(idx)))
    fig, ax = plt.subplots(figsize=(10, max(6, len(idx) * 0.22)))
    ax.barh([names[i] for i in idx], [imp[i] for i in idx], color=colors)
    ax.set_xlabel("Mean Gini Decrease (normalised, 5-fold CV)")
    ax.set_ylabel("Gene")
    ax.set_title("DT Feature Importance — %s vs. GTEx\n(top-%d genes, 5-fold CV mean)" % (label, len(genes)))
    plt.tight_layout()
    plt.savefig(out_path, dpi=150, bbox_inches="tight")
    plt.close()

def plot_cm(cm_arr, cohort, acc, f1mac, out_path):
    fig, ax = plt.subplots(figsize=(5, 4))
    sns.heatmap(cm_arr, annot=True, fmt="d", cmap="Blues", ax=ax,
                xticklabels=["GTEx Healthy", "TCGA Tumor"],
                yticklabels=["GTEx Healthy", "TCGA Tumor"])
    ax.set_xlabel("Predicted"); ax.set_ylabel("True")
    ax.set_title("DT — %s vs GTEx\nAcc=%.4f  MacroF1=%.4f" % (cohort, acc, f1mac))
    plt.tight_layout()
    plt.savefig(out_path, dpi=150, bbox_inches="tight")
    plt.close()

def run_cohort(cohort, sym_map):
    vs_path = DS_DIR + "/dataset_%s_vs_gtex.csv" % cohort
    dt_path = OUT_DIR + "/Biomarkers_DT_%s.csv" % cohort
    label   = COHORT_LABELS.get(cohort, cohort)

    if not os.path.exists(vs_path) or not os.path.exists(dt_path):
        print("  SKIP: missing files")
        return None

    df = pd.read_csv(vs_path)
    if df["target"].nunique() < 2:
        print("  SKIP: single-class")
        return None

    genes = [g for g in pd.read_csv(dt_path).head(TOP_N)["Gene"].tolist() if g in df.columns]
    if not genes:
        print("  SKIP: no genes")
        return None
    print("  Genes: %d" % len(genes))

    X = np.log2(df[genes].apply(pd.to_numeric, errors="coerce").fillna(0) + 1.0).values
    y = df["target"].astype(int).values

    skf = StratifiedKFold(n_splits=FOLDS, shuffle=True, random_state=SEED)
    fold_imps, all_preds, fold_accs, fold_f1s = [], np.zeros(len(y), dtype=int), [], []

    for tr, te in skf.split(X, y):
        clf = DecisionTreeClassifier(max_depth=None, min_samples_leaf=10,
                                     class_weight="balanced", random_state=SEED)
        clf.fit(X[tr], y[tr])
        preds = clf.predict(X[te])
        all_preds[te] = preds
        fold_imps.append(clf.feature_importances_)
        fold_accs.append(float((preds == y[te]).mean()))
        fold_f1s.append(float(f1_score(y[te], preds, average="macro", zero_division=0)))

    mean_imp = np.mean(fold_imps, axis=0)
    mean_imp = mean_imp / (mean_imp.max() + 1e-10)
    acc   = float(np.mean(fold_accs))
    f1mac = float(np.mean(fold_f1s))
    cm_arr = confusion_matrix(y, all_preds)

    fi_path = FIG_DIR + "/feature_importance_DT_%s_vs_gtex.png" % cohort
    cm_path = OUT_DIR + "/confusion_matrix_DT_%s_vs_GTEx.png" % cohort
    kp_path = FIG_DIR + "/pathway_DT_%s_vs_gtex.png" % cohort

    plot_importance(mean_imp, genes, sym_map, label, fi_path)
    plot_cm(cm_arr, cohort, acc, f1mac, cm_path)

    top_idx   = np.argsort(mean_imp)[::-1][:20]
    top_genes = [sym_map.get(genes[i], genes[i]) for i in top_idx]
    kegg_rows = enrichr_kegg(top_genes, label)
    kegg_ok   = plot_kegg(kegg_rows, cohort, kp_path)
    print("  Acc=%.4f F1=%.4f KEGG=%s" % (acc, f1mac, "OK" if kegg_ok else "FAIL"))

    return {"cohort": cohort, "label": label, "tissue": TISSUE_MAP.get(cohort,""),
            "n_genes": len(genes), "acc": round(acc,4), "f1_macro": round(f1mac,4),
            "fi_path": fi_path, "cm_path": cm_path,
            "kp_path": kp_path if kegg_ok else None}

def embed_in_docx(results):
    doc = docx.Document(DOCX_PATH)
    p   = doc.paragraphs
    img_paras = [i for i,para in enumerate(p)
                 if para._element.findall('.//' + qn("a:blip"), para._element.nsmap)]
    next_fig = len(img_paras) + 1
    print("Existing figures: %d — new start at Fig %d" % (len(img_paras), next_fig))

    def add_bold(text):
        hp = doc.add_paragraph(); hp.add_run(text).bold = True

    def add_fig(img_path, cap):
        if not os.path.exists(img_path):
            print("  MISSING:", img_path); return
        p_img = doc.add_paragraph()
        p_img.add_run().add_picture(img_path, width=Inches(5.5))
        cap_p = doc.add_paragraph(cap)
        try: cap_p.style = doc.styles["Caption"]
        except: pass

    add_bold("V. Decision Tree Validation on GTEx Healthy Tissue")
    doc.add_paragraph(
        "Each cohort-specific DT model (top-50 DT biomarkers, max_depth=None, "
        "min_samples_leaf=10, 5-fold StratifiedKFold) was applied to GTEx healthy samples "
        "to validate tumour-specificity of the identified biomarkers."
    )

    fig_n = next_fig
    for res in results:
        if res is None: continue
        add_bold("%s vs. GTEx (%s)" % (res["cohort"], res["tissue"]))

        add_fig(res["fi_path"],
                "Figure %d. DT feature importance for %s vs. GTEx "
                "(top-%d genes, 5-fold CV mean Gini decrease, normalised)." % (fig_n, res["label"], TOP_N))
        fig_n += 1

        if res["kp_path"]:
            add_fig(res["kp_path"],
                    "Figure %d. KEGG pathway enrichment (Enrichr, KEGG_2021_Human) "
                    "for top-20 DT biomarkers of %s vs. GTEx." % (fig_n, res["label"]))
            fig_n += 1

        add_fig(res["cm_path"],
                "Figure %d. OOF confusion matrix for DT on %s vs. GTEx "
                "(5-fold StratifiedKFold; Acc=%.4f, Macro-F1=%.4f)." % (
                    fig_n, res["label"], res["acc"], res["f1_macro"]))
        fig_n += 1

        doc.add_paragraph(
            "%s vs. GTEx (%s): Accuracy=%.4f, Macro-F1=%.4f (%d/%d DT genes present)." % (
                res["cohort"], res["tissue"], res["acc"], res["f1_macro"],
                res["n_genes"], TOP_N))

    doc.save(DOCX_PATH)
    print("Docx saved. New figures: %d" % (fig_n - next_fig))

def save_summary(results):
    out = OUT_DIR + "/DT_GTEx_Summary.csv"
    with open(out, "w", newline="") as f:
        w = csv.DictWriter(f, fieldnames=["cohort","tissue","n_genes","acc","f1_macro"])
        w.writeheader()
        for r in results:
            if r: w.writerow({k: r[k] for k in ["cohort","tissue","n_genes","acc","f1_macro"]})
    print("Summary:", out)

def main():
    os.makedirs(FIG_DIR, exist_ok=True)
    sym_map = load_symbol_map()
    print("Symbol map: %d entries" % len(sym_map))
    vs_files = sorted(glob.glob(DS_DIR + "/dataset_TCGA_*_vs_gtex.csv"))
    cohorts  = [os.path.basename(f).replace("dataset_","").replace("_vs_gtex.csv","")
                for f in vs_files]
    print("Cohorts:", cohorts)
    results = []
    for cohort in cohorts:
        print("\n=== %s ===" % cohort)
        results.append(run_cohort(cohort, sym_map))
    save_summary(results)
    embed_in_docx(results)
    print("\nAll done.")

if __name__ == "__main__":
    main()
